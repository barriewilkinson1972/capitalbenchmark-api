from __future__ import annotations

import os
import re
import shutil
import subprocess
import tempfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

try:
    from credit_memo import format_currency, format_multiple, format_percent
except Exception:  # pragma: no cover - useful if rendering from standalone payloads
    def _safe_float(value: Any) -> float | None:
        try:
            if value is None:
                return None
            value = float(value)
            if value != value or value in (float("inf"), float("-inf")):
                return None
            return value
        except Exception:
            return None

    def format_currency(value: Any) -> str:
        number = _safe_float(value)
        if number is None:
            return "n/a"
        sign = "-" if number < 0 else ""
        number = abs(number)
        if number >= 1_000_000_000_000:
            return f"{sign}${number / 1_000_000_000_000:.2f}tn"
        if number >= 1_000_000_000:
            return f"{sign}${number / 1_000_000_000:.1f}bn"
        if number >= 1_000_000:
            return f"{sign}${number / 1_000_000:.1f}mn"
        return f"{sign}${number:,.0f}"

    def format_percent(value: Any) -> str:
        number = _safe_float(value)
        return "n/a" if number is None else f"{number * 100:.2f}%"

    def format_multiple(value: Any) -> str:
        number = _safe_float(value)
        return "n/a" if number is None else f"{number:.1f}x"


DOCX_MIMETYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
PDF_MIMETYPE = "application/pdf"

BRAND_NAVY = "0B1F3A"
BRAND_BLUE = "1F4E79"
LIGHT_BLUE = "D9EAF7"
LIGHT_GREY = "F3F6F8"
BORDER_GREY = "D9E2EA"
WHITE = "FFFFFF"
DARK_GREY = "555555"


def _safe_text(value: Any, default: str = "") -> str:
    if value is None:
        return default
    text = str(value).strip()
    if not text or text.lower() in {"nan", "none", "null"}:
        return default
    return text


def _short(value: Any, max_chars: int = 550) -> str:
    text = _safe_text(value)
    if len(text) <= max_chars:
        return text
    return text[: max_chars - 3].rstrip() + "..."


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_border(cell, color: str = BORDER_GREY, size: str = "4") -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def _set_cell_text(
    cell,
    text: str,
    bold: bool = False,
    font_size: float = 8.2,
    color: str | None = None,
    align: str = "left",
) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    if align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(_safe_text(text))
    run.bold = bold
    run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0


def _setup_styles(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(9.0)
    styles["Heading 1"].font.name = "Aptos Display"
    styles["Heading 1"].font.size = Pt(14)
    styles["Heading 1"].font.bold = True
    styles["Heading 1"].font.color.rgb = RGBColor.from_string(BRAND_NAVY)
    styles["Heading 2"].font.name = "Aptos Display"
    styles["Heading 2"].font.size = Pt(11)
    styles["Heading 2"].font.bold = True
    styles["Heading 2"].font.color.rgb = RGBColor.from_string(BRAND_BLUE)


def _add_footer(doc: Document) -> None:
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Capital Benchmark - controlled first-draft credit memo")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string("666666")


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(8 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(4)


def _add_paragraph(doc: Document, text: str | None) -> None:
    text = _safe_text(text)
    if not text:
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.05
    p.add_run(text)


def _add_bullets(doc: Document, items: list[Any] | None) -> None:
    items = items or []
    if not items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.add_run("n/a")
        return
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.add_run(_safe_text(item))


def _add_key_value_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    _set_cell_text(hdr[0], "Field", bold=True, color=WHITE, font_size=8.2)
    _set_cell_text(hdr[1], "Value", bold=True, color=WHITE, font_size=8.2)
    for cell in hdr:
        _set_cell_shading(cell, BRAND_NAVY)
        _set_cell_border(cell)
    for label, value in rows:
        cells = table.add_row().cells
        _set_cell_text(cells[0], label, bold=True, font_size=8.2)
        _set_cell_text(cells[1], value, font_size=8.2)
        for cell in cells:
            _set_cell_border(cell)
    doc.add_paragraph()


def _add_data_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    font_size: float = 7.3,
    right_align_cols: set[int] | None = None,
) -> None:
    right_align_cols = right_align_cols or set()
    if not rows:
        rows = [["n/a"] + [""] * (len(headers) - 1)]
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    table.style = "Table Grid"
    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        _set_cell_text(cell, header, bold=True, font_size=font_size, color=WHITE, align="center")
        _set_cell_shading(cell, BRAND_NAVY)
        _set_cell_border(cell)
    for i, row in enumerate(rows):
        cells = table.add_row().cells
        fill = LIGHT_GREY if i % 2 else WHITE
        for j, value in enumerate(row):
            align = "right" if j in right_align_cols else "left"
            _set_cell_text(cells[j], _safe_text(value), font_size=font_size, align=align)
            _set_cell_shading(cells[j], fill)
            _set_cell_border(cells[j])
    doc.add_paragraph()


def _request_rows(context: dict[str, Any]) -> list[tuple[str, str]]:
    request = context.get("credit_request", {})
    exposure = context.get("exposure_analytics", {})
    tenor = request.get("tenor_years")
    return [
        ("Request type", _safe_text(request.get("request_type"), "n/a")),
        ("Facility type", _safe_text(request.get("facility_type"), "n/a")),
        ("Purpose", _safe_text(request.get("purpose"), "n/a")),
        ("Existing exposure", format_currency(request.get("existing_exposure_usd"))),
        ("Requested increase", format_currency(request.get("requested_increase_usd"))),
        ("Proposed exposure", format_currency(request.get("proposed_exposure_usd"))),
        ("Tenor", f"{tenor:g} years" if isinstance(tenor, (int, float)) else "n/a"),
        ("Secured", _safe_text(request.get("secured"), "n/a")),
        ("Base EL on proposed exposure", format_currency(exposure.get("base_expected_loss_on_proposed_exposure"))),
    ]


def _rating_rows(context: dict[str, Any]) -> list[tuple[str, str]]:
    rating = context.get("capital_benchmark_rating", {})
    raw = rating.get("cb_rating_num_raw")
    raw_text = f"{raw:.2f}" if isinstance(raw, (int, float)) else "n/a"
    return [
        ("CB rating", _safe_text(rating.get("cb_rating"), "n/a")),
        ("CB PD", format_percent(rating.get("cb_pd"))),
        ("Raw model score", raw_text),
        ("Agency rating", _safe_text(rating.get("agency_rating"), "n/a")),
        ("CB vs agency", _safe_text(rating.get("cb_vs_agency_comment"), "n/a")),
        ("Rating context quality", _safe_text(rating.get("rating_context_quality"), "n/a")),
        ("Model", _safe_text(rating.get("model_type"), "n/a")),
        ("Model version", _safe_text(rating.get("model_version"), "n/a")),
    ]


def _financial_rows(context: dict[str, Any]) -> list[tuple[str, str]]:
    financials = context.get("financials", {})
    metrics = context.get("credit_metrics", {})
    return [
        ("Revenue", format_currency(financials.get("total_revenue_usd"))),
        ("Total debt", format_currency(financials.get("total_debt_usd"))),
        ("Cash", format_currency(financials.get("total_cash_usd"))),
        ("Net debt", format_currency(financials.get("net_debt_usd"))),
        ("EBITDA", format_currency(financials.get("ebitda_usd"))),
        ("Market cap", format_currency(financials.get("market_cap_usd"))),
        ("Enterprise value", format_currency(financials.get("enterprise_value_usd"))),
        ("Debt / revenue", format_multiple(metrics.get("debt_to_revenue"))),
        ("Debt / EBITDA", format_multiple(metrics.get("debt_to_ebitda"))),
        ("Net debt / EBITDA", format_multiple(metrics.get("net_debt_to_ebitda"))),
        ("Cash / debt", format_percent(metrics.get("cash_to_debt"))),
        ("Profit margin", format_percent(metrics.get("profit_margin"))),
        ("Annual volatility", format_percent(metrics.get("annual_vol_5y"))),
    ]


def _driver_rows(context: dict[str, Any]) -> list[list[str]]:
    rows = []
    for item in context.get("rating_driver_sensitivities", [])[:8]:
        effect = item.get("rating_effect_vs_median")
        rows.append([
            _safe_text(item.get("label"), "n/a"),
            f"{effect:+.2f} notches" if isinstance(effect, (int, float)) else "n/a",
            _safe_text(item.get("interpretation"), "n/a"),
        ])
    return rows


def _watchpoint_rows(context: dict[str, Any]) -> list[list[str]]:
    rows = []
    for item in context.get("financial_watchpoints", []):
        rows.append([
            _safe_text(item.get("metric"), "n/a"),
            _safe_text(item.get("value"), "n/a"),
            _safe_text(item.get("peer_value"), "n/a"),
            _safe_text(item.get("severity"), "n/a"),
            _safe_text(item.get("comment"), "n/a"),
        ])
    return rows


def _policy_rows(context: dict[str, Any]) -> list[list[str]]:
    rows = []
    for item in context.get("policy_evaluation", {}).get("triggered_policies", []):
        rows.append([
            _safe_text(item.get("policy_id"), "n/a"),
            _safe_text(item.get("name"), "n/a"),
            _safe_text(item.get("observed_text"), "n/a"),
            _safe_text(item.get("severity"), "n/a"),
            _safe_text(item.get("required_action"), "n/a"),
        ])
    return rows


def _peer_rows(context: dict[str, Any]) -> list[tuple[str, str]]:
    peer = context.get("peer_and_anchor_context", {})
    return [
        ("Industry rating anchor", _safe_text(peer.get("industry_rating_anchor"), "n/a")),
        ("Country rating anchor", _safe_text(peer.get("country_rating_anchor"), "n/a")),
        ("Industry median debt / revenue", format_multiple(peer.get("industry_median_debt_to_revenue"))),
        ("Industry median profit margin", format_percent(peer.get("industry_median_profit_margin"))),
        ("Industry median annual volatility", format_percent(peer.get("industry_median_annual_vol_5y"))),
        ("Industry median revenue", format_currency(peer.get("industry_median_revenue_usd"))),
        ("Revenue percentile - industry", format_percent(peer.get("revenue_percentile_industry"))),
        ("Profit margin percentile - industry", format_percent(peer.get("profit_margin_percentile_industry"))),
        ("Relative debt percentile - industry", format_percent(peer.get("relative_debt_percentile_industry"))),
        ("Volatility percentile - industry", format_percent(peer.get("volatility_percentile_industry"))),
    ]


def _experiment_rows(payload: dict[str, Any]) -> list[tuple[str, str]]:
    config = payload.get("experiment_config") or {}
    if not config:
        return []
    return [
        ("Experiment ID", _safe_text(config.get("experiment_id"), "n/a")),
        ("Context mode", _safe_text(config.get("context_mode"), "n/a")),
        ("Policy mode", _safe_text(config.get("policy_mode"), "n/a")),
        ("Prompt mode", _safe_text(config.get("prompt_mode"), "n/a")),
        ("Model", _safe_text(config.get("model"), "n/a")),
        ("LLM sees deterministic context", _safe_text(config.get("llm_sees_full_deterministic_context"), "n/a")),
        ("LLM sees credit policy", _safe_text(config.get("llm_sees_credit_policy"), "n/a")),
        ("LLM sees policy evaluation", _safe_text(config.get("llm_sees_policy_evaluation"), "n/a")),
    ]


def _add_preamble(doc: Document, narrative: dict[str, Any], key: str) -> None:
    preambles = narrative.get("section_preambles") or {}
    if not isinstance(preambles, dict):
        return
    text = _safe_text(preambles.get(key))
    if not text:
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.05
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = RGBColor.from_string(DARK_GREY)


def render_credit_memo_docx(
    memo_payload: dict[str, Any],
    output_path: str | Path,
) -> Path:
    """Render the payload returned by create_credit_memo() to a DOCX credit memo."""
    output_path = Path(output_path)
    narrative = memo_payload.get("narrative", {})
    context = memo_payload.get("memo_context", {})
    borrower = context.get("borrower", {})

    doc = Document()
    _setup_styles(doc)
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.50)
    section.right_margin = Inches(0.50)
    _add_footer(doc)

    title = narrative.get("title") or "Capital Benchmark Credit Memo"
    borrower_name = narrative.get("borrower_name") or borrower.get("company_name") or borrower.get("symbol") or "Borrower"
    generated_at = context.get("generated_at_utc") or datetime.now(timezone.utc).isoformat()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(_safe_text(title))
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor.from_string(BRAND_NAVY)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(
        f"Borrower: {borrower_name}\n"
        f"Ticker: {_safe_text(borrower.get('symbol'), 'n/a')} | "
        f"Sector: {_safe_text(borrower.get('sector'), 'n/a')} | "
        f"Country: {_safe_text(borrower.get('country'), 'n/a')}\n"
        f"Generated: {generated_at}"
    )
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(DARK_GREY)

    experiment_rows = _experiment_rows(memo_payload)
    if experiment_rows:
        _add_heading(doc, "Generation Configuration", 1)
        _add_key_value_table(doc, experiment_rows)

    _add_heading(doc, "Introduction", 1)
    _add_paragraph(doc, narrative.get("introduction"))

    _add_heading(doc, "Executive Summary", 1)
    _add_paragraph(doc, narrative.get("executive_summary"))

    _add_heading(doc, "Borrower and Request Summary", 1)
    _add_paragraph(doc, narrative.get("request_summary"))
    _add_key_value_table(doc, _request_rows(context))

    _add_heading(doc, "Business Profile", 1)
    _add_paragraph(doc, narrative.get("business_profile") or _short(borrower.get("business_summary"), 1000))

    _add_heading(doc, "Capital Benchmark Rating Assessment", 1)
    _add_paragraph(doc, narrative.get("capital_benchmark_rating_assessment"))
    _add_key_value_table(doc, _rating_rows(context))

    _add_heading(doc, "Rating Driver Commentary", 1)
    _add_paragraph(doc, narrative.get("rating_driver_commentary"))
    _add_data_table(
        doc,
        ["Driver", "Effect vs median", "Assessment"],
        _driver_rows(context),
        font_size=7.2,
        right_align_cols={1},
    )

    _add_heading(doc, "Positive Rating Drivers", 2)
    _add_bullets(doc, narrative.get("positive_rating_drivers", []))
    _add_heading(doc, "Negative Rating Drivers", 2)
    _add_bullets(doc, narrative.get("negative_rating_drivers", []))
    _add_heading(doc, "Neutral Rating Diagnostics", 2)
    _add_bullets(doc, narrative.get("neutral_rating_diagnostics", []))

    _add_heading(doc, "Financial Risk Assessment", 1)
    _add_paragraph(doc, narrative.get("financial_risk_assessment"))
    _add_key_value_table(doc, _financial_rows(context))

    _add_heading(doc, "Financial Watchpoints", 2)
    _add_data_table(
        doc,
        ["Metric", "Value", "Peer", "Severity", "Comment"],
        _watchpoint_rows(context),
        font_size=6.8,
        right_align_cols={1, 2},
    )

    _add_heading(doc, "Credit Policy Compliance", 1)
    _add_preamble(doc, narrative, "policy_compliance")
    _add_paragraph(doc, narrative.get("policy_compliance_assessment"))
    _add_data_table(
        doc,
        ["Policy ID", "Rule", "Observed", "Severity", "Required action"],
        _policy_rows(context),
        font_size=6.5,
    )
    _add_heading(doc, "Policy Breaches and Triggers", 2)
    _add_bullets(doc, narrative.get("policy_breaches", []))
    _add_heading(doc, "Required Policy Actions", 2)
    _add_bullets(doc, narrative.get("policy_required_actions", []))
    _add_heading(doc, "Policy Missing Information", 2)
    _add_bullets(doc, narrative.get("policy_missing_information", []))
    _add_heading(doc, "Policy Escalation Assessment", 2)
    _add_paragraph(doc, narrative.get("policy_escalation_assessment"))

    _add_heading(doc, "Peer and Anchor Context", 1)
    _add_paragraph(doc, narrative.get("peer_and_anchor_context"))
    _add_key_value_table(doc, _peer_rows(context))

    _add_heading(doc, "Key Credit Strengths", 1)
    _add_bullets(doc, narrative.get("key_credit_strengths", []))

    _add_heading(doc, "Key Credit Watchpoints", 1)
    _add_bullets(doc, narrative.get("key_credit_watchpoints", []))

    _add_heading(doc, "Questions for Relationship Manager", 1)
    _add_bullets(doc, narrative.get("questions_for_relationship_manager", []))

    _add_heading(doc, "Credit Committee Focus Areas", 1)
    _add_bullets(doc, narrative.get("credit_committee_focus_areas", []))

    _add_heading(doc, "Conclusion and Recommendation", 1)
    _add_paragraph(doc, narrative.get("conclusion_recommendation"))

    _add_heading(doc, "Data Quality and Limitations", 1)
    _add_paragraph(doc, narrative.get("data_quality_and_limitations"))
    warnings = context.get("data_quality", {}).get("warnings", [])
    if warnings:
        _add_heading(doc, "Data Quality Warnings", 2)
        _add_bullets(doc, warnings)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path


def convert_docx_to_pdf(docx_path: str | Path, output_dir: str | Path | None = None) -> Path:
    """
    Convert a DOCX to PDF using LibreOffice headless.

    Production note: systemd often sets PATH only to the Python venv, so prefer
    LIBREOFFICE_BIN when provided and explicitly add /usr/bin to PATH.
    """
    docx_path = Path(docx_path).resolve()
    output_dir = Path(output_dir).resolve() if output_dir is not None else docx_path.parent
    output_dir.mkdir(parents=True, exist_ok=True)

    libreoffice = (
        os.getenv("LIBREOFFICE_BIN")
        or shutil.which("libreoffice")
        or shutil.which("soffice")
        or "/usr/bin/libreoffice"
    )

    if not Path(libreoffice).exists() and shutil.which(libreoffice) is None:
        raise RuntimeError(
            "LibreOffice is not installed or not accessible. "
            f"Tried: {libreoffice}. On Ubuntu install with: sudo apt install -y libreoffice"
        )

    with tempfile.TemporaryDirectory() as profile_dir:
        profile_uri = Path(profile_dir).resolve().as_uri()
        cmd = [
            libreoffice,
            "--headless",
            "--nologo",
            "--nofirststartwizard",
            "--nolockcheck",
            f"-env:UserInstallation={profile_uri}",
            "--convert-to",
            "pdf",
            "--outdir",
            str(output_dir),
            str(docx_path),
        ]
        env = os.environ.copy()
        env["HOME"] = env.get("HOME") or "/tmp"
        env["PATH"] = "/usr/local/sbin:/usr/local/bin:/usr/sbin:/usr/bin:/sbin:/bin:" + env.get("PATH", "")
        completed = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            timeout=90,
            env=env,
        )
        if completed.returncode != 0:
            raise RuntimeError(
                "LibreOffice PDF conversion failed.\n"
                f"Command: {' '.join(cmd)}\n"
                f"Return code: {completed.returncode}\n"
                f"stdout: {completed.stdout}\n"
                f"stderr: {completed.stderr}"
            )

    pdf_path = output_dir / f"{docx_path.stem}.pdf"
    if not pdf_path.exists() or pdf_path.stat().st_size == 0:
        raise RuntimeError(
            "LibreOffice did not create the expected PDF file.\n"
            f"Expected: {pdf_path}\n"
            f"Output directory contents: {[p.name for p in output_dir.iterdir()]}"
        )
    return pdf_path


def credit_memo_filename(
    borrower: dict[str, Any] | None = None,
    extension: str = "docx",
    prefix: str = "capital_benchmark_credit_memo",
) -> str:
    borrower = borrower or {}
    symbol = _safe_text(borrower.get("symbol"), "borrower")
    company = _safe_text(borrower.get("company_name"), "")
    stamp = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
    raw = f"{prefix}_{symbol}_{company}_{stamp}.{extension}"
    raw = raw.replace("+", "plus").replace("-", "minus")
    return re.sub(r"[^A-Za-z0-9_.-]+", "_", raw).strip("_")

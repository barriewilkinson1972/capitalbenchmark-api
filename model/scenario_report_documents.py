from __future__ import annotations

import os
import re
import shutil
import subprocess
import tempfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from model.scenario_report import format_currency, format_multiple, format_percent

DOCX_MIMETYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
PDF_MIMETYPE = "application/pdf"

BRAND_NAVY = "0B1F3A"
BRAND_BLUE = "1F4E79"
LIGHT_BLUE = "D9EAF7"
LIGHT_GREY = "F3F6F8"
BORDER_GREY = "D9E2EA"
WHITE = "FFFFFF"


def _safe_text(value: Any) -> str:
    if value is None:
        return ""
    return str(value)


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_border(cell, color: str = BORDER_GREY, size: str = "4") -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def _set_cell_text(
    cell,
    text: str,
    bold: bool = False,
    font_size: float = 8.5,
    color: str | None = None,
    align: str = "left",
) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    if align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0


def _add_paragraph(doc: Document, text: str, style: str | None = None) -> None:
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.05
    p.add_run(text)


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(8 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(4)


def _add_bullets(doc: Document, items: list[Any]) -> None:
    for item in items or []:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.add_run(_safe_text(item))


def _add_key_value_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    _set_cell_text(hdr[0], "Metric", bold=True, color=WHITE, font_size=8.5)
    _set_cell_text(hdr[1], "Value", bold=True, color=WHITE, font_size=8.5)
    for cell in hdr:
        _set_cell_shading(cell, BRAND_NAVY)
        _set_cell_border(cell)
    for label, value in rows:
        cells = table.add_row().cells
        _set_cell_text(cells[0], label, bold=True, font_size=8.5)
        _set_cell_text(cells[1], value, font_size=8.5)
        for cell in cells:
            _set_cell_border(cell)
    doc.add_paragraph()


def _add_data_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    font_size: float = 7.5,
    right_align_cols: set[int] | None = None,
) -> None:
    right_align_cols = right_align_cols or set()
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    table.style = "Table Grid"

    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        _set_cell_text(cell, header, bold=True, font_size=font_size, color=WHITE, align="center")
        _set_cell_shading(cell, BRAND_NAVY)
        _set_cell_border(cell)

    for i, row in enumerate(rows):
        cells = table.add_row().cells
        fill = LIGHT_GREY if i % 2 else WHITE
        for j, value in enumerate(row):
            align = "right" if j in right_align_cols else "left"
            _set_cell_text(cells[j], value, font_size=font_size, align=align)
            _set_cell_shading(cells[j], fill)
            _set_cell_border(cells[j])
    doc.add_paragraph()


def _portfolio_rows(report_context: dict[str, Any]) -> list[tuple[str, str]]:
    summary = report_context.get("portfolio_summary", {})
    rows = [
        ("Total EAD", format_currency(summary.get("total_ead"))),
        ("Weighted base PD", format_percent(summary.get("weighted_base_pd"))),
        ("Weighted stressed PD", format_percent(summary.get("weighted_stressed_pd"))),
        ("Base expected loss", format_currency(summary.get("base_expected_loss"))),
        ("Stressed expected loss", format_currency(summary.get("stressed_expected_loss"))),
        ("Stressed expected loss / EAD", format_percent(summary.get("stressed_loss_rate"))),
        ("Expected loss multiple", format_multiple(summary.get("expected_loss_multiple"))),
    ]
    if summary.get("scenario_tail_probability") is not None:
        rows.append(("Scenario tail probability", format_percent(summary.get("scenario_tail_probability"))))
    if summary.get("scenario_tail_odds") is not None:
        rows.append(("Scenario tail odds", f"1 in {summary.get('scenario_tail_odds'):,.0f}"))
    return rows


def _industry_rows(report_context: dict[str, Any], limit: int = 10) -> list[list[str]]:
    rows = []
    for row in report_context.get("top_industries", [])[:limit]:
        rows.append([
            _safe_text(row.get("industry")),
            _safe_text(row.get("obligors")),
            format_currency(row.get("ead")),
            format_percent(row.get("stressed_pd")),
            format_currency(row.get("stressed_expected_loss")),
            format_multiple(row.get("pd_multiple")),
        ])
    return rows


def _obligor_rows(report_context: dict[str, Any], limit: int = 10) -> list[list[str]]:
    rows = []
    for row in report_context.get("top_obligors", [])[:limit]:
        rows.append([
            _safe_text(row.get("company_name") or row.get("symbol")),
            _safe_text(row.get("industry")),
            _safe_text(row.get("agency_rating") or row.get("cb_rating")),
            format_currency(row.get("ead")),
            format_percent(row.get("stressed_pd")),
            format_currency(row.get("stressed_expected_loss")),
        ])
    return rows


def _setup_styles(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(9.0)
    styles["Heading 1"].font.name = "Aptos Display"
    styles["Heading 1"].font.size = Pt(14)
    styles["Heading 1"].font.bold = True
    styles["Heading 1"].font.color.rgb = RGBColor.from_string(BRAND_NAVY)
    styles["Heading 2"].font.name = "Aptos Display"
    styles["Heading 2"].font.size = Pt(11)
    styles["Heading 2"].font.bold = True
    styles["Heading 2"].font.color.rgb = RGBColor.from_string(BRAND_BLUE)


def _add_footer(doc: Document) -> None:
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Capital Benchmark - Scenario analytics generated from model output")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string("666666")


def render_scenario_report_docx(
    report_payload: dict[str, Any],
    output_path: str | Path,
    industry_limit: int = 10,
    obligor_limit: int = 10,
) -> Path:
    """Render the scenario report payload returned by create_scenario_report() to DOCX."""
    output_path = Path(output_path)
    narrative = report_payload.get("narrative", {})
    report_context = report_payload.get("report_context", {})

    doc = Document()
    _setup_styles(doc)
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.50)
    section.right_margin = Inches(0.50)
    _add_footer(doc)

    title = narrative.get("title") or "Capital Benchmark Scenario Report"
    scenario_label = narrative.get("scenario_label") or "Custom stress scenario"
    generated_at = report_context.get("generated_at_utc") or datetime.now(timezone.utc).isoformat()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor.from_string(BRAND_NAVY)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(f"Scenario: {scenario_label}\nGenerated: {generated_at}")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string("555555")

    _add_heading(doc, "Executive Summary", 1)
    _add_paragraph(doc, narrative.get("executive_summary", ""))

    _add_heading(doc, "Scenario Interpretation", 1)
    _add_paragraph(doc, narrative.get("scenario_interpretation", ""))

    _add_heading(doc, "Portfolio Impact", 1)
    _add_paragraph(doc, narrative.get("portfolio_impact", ""))
    _add_key_value_table(doc, _portfolio_rows(report_context))

    _add_heading(doc, "Industry Concentration", 1)
    _add_paragraph(doc, narrative.get("industry_concentration", ""))
    _add_data_table(
        doc,
        ["Industry", "Obligors", "EAD", "Stressed PD", "Stressed EL", "PD Multiple"],
        _industry_rows(report_context, industry_limit),
        font_size=7.0,
        right_align_cols={1, 2, 3, 4, 5},
    )

    _add_heading(doc, "Obligor Concentration", 1)
    _add_paragraph(doc, narrative.get("obligor_concentration", ""))
    _add_data_table(
        doc,
        ["Obligor", "Industry", "Rating", "EAD", "Stressed PD", "Stressed EL"],
        _obligor_rows(report_context, obligor_limit),
        font_size=6.8,
        right_align_cols={3, 4, 5},
    )

    _add_heading(doc, "Key Risk Drivers", 1)
    _add_bullets(doc, narrative.get("key_risk_drivers", []))

    _add_heading(doc, "Management Takeaways", 1)
    _add_bullets(doc, narrative.get("management_takeaways", []))

    _add_heading(doc, "Limitations", 1)
    _add_paragraph(doc, narrative.get("limitations", ""))

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path


def convert_docx_to_pdf(docx_path: str | Path, output_dir: str | Path | None = None) -> Path:
    """Convert a DOCX to PDF using LibreOffice headless. Requires libreoffice on PATH."""
    docx_path = Path(docx_path)
    output_dir = Path(output_dir) if output_dir is not None else docx_path.parent
    output_dir.mkdir(parents=True, exist_ok=True)

    libreoffice = shutil.which("libreoffice") or shutil.which("soffice")
    if not libreoffice:
        raise RuntimeError(
            "LibreOffice is not installed or not on PATH. Install it with: sudo apt install -y libreoffice"
        )

    with tempfile.TemporaryDirectory() as profile_dir:
        profile_uri = Path(profile_dir).resolve().as_uri()
        cmd = [
            libreoffice,
            "--headless",
            "--nologo",
            "--nofirststartwizard",
            f"-env:UserInstallation={profile_uri}",
            "--convert-to",
            "pdf",
            "--outdir",
            str(output_dir),
            str(docx_path),
        ]
        completed = subprocess.run(cmd, capture_output=True, text=True, timeout=90)
        if completed.returncode != 0:
            raise RuntimeError(
                "LibreOffice PDF conversion failed.\n"
                f"stdout: {completed.stdout}\n"
                f"stderr: {completed.stderr}"
            )

    pdf_path = output_dir / f"{docx_path.stem}.pdf"
    if not pdf_path.exists() or pdf_path.stat().st_size == 0:
        raise RuntimeError("LibreOffice did not create the expected PDF file.")
    return pdf_path


def scenario_report_filename(
    scenario: dict[str, Any] | None = None,
    extension: str = "docx",
    prefix: str = "capital_benchmark_scenario_report",
) -> str:
    scenario = scenario or {}
    market = scenario.get("market", "")
    technology = scenario.get("technology", "")
    commodity = scenario.get("commodity", "")
    stamp = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
    raw = f"{prefix}_m{market}_t{technology}_c{commodity}_{stamp}.{extension}"
    return re.sub(r"[^A-Za-z0-9_.-]+", "_", raw).replace("+", "plus").replace("-", "minus")
